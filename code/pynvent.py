
"""
Inventory Document Generator
============================

Este módulo permite generar un documento de inventario en formato .docx a partir de un directorio de imágenes
ordenadas en carpetas. La estructura del documento respeta la jerarquía de carpetas, usando los nombres de 
las carpetas como títulos y listando las imágenes con su correspondiente miniatura dentro del archivo. 
Además, incluye la opción de utilizar el reconocimiento de objetos mediante la API de Google Vision para 
etiquetar y renombrar automáticamente las imágenes según su contenido antes de generar el documento.

Funcionalidades principales:
- Recorrido recursivo de un directorio para estructurar el inventario de imágenes por secciones y subsecciones.
- Renombrado de imágenes mediante etiquetas obtenidas con Google Vision (DE PAGO).
- Generación de un documento de inventario en formato .docx que incluye imágenes redimensionadas para 
  ajustarse a los márgenes del documento.

Requisitos:
- python-docx
- pillow
- google-cloud-vision

Autor:
- Alba GC

Fecha:
- 2024

Copyright (c) 2024 Alba

Licencia MIT:
------------
Se concede permiso, sin cargo, para cualquier persona que obtenga una copia de este software y de los archivos
de documentación asociados (el "Software"), para tratar el Software sin restricción, incluyendo sin limitación 
los derechos de usar, copiar, modificar, fusionar, publicar, distribuir, sublicenciar y/o vender copias del 
Software, y permitir a las personas a las que se les proporcione el Software que lo hagan, sujeto a las 
siguientes condiciones:

El aviso de copyright anterior y este aviso de permiso se incluirán en todas las copias o partes sustanciales 
del Software.

EL SOFTWARE SE PROPORCIONA "TAL CUAL", SIN GARANTÍA DE NINGÚN TIPO, EXPRESA O IMPLÍCITA, INCLUYENDO PERO NO 
LIMITADO A GARANTÍAS DE COMERCIALIZACIÓN, IDONEIDAD PARA UN PROPÓSITO PARTICULAR E INCUMPLIMIENTO. EN NINGÚN 
CASO LOS AUTORES O TITULARES DEL COPYRIGHT SERÁN RESPONSABLES DE NINGUNA RECLAMACIÓN, DAÑO U OTRA 
RESPONSABILIDAD, YA SEA EN UNA ACCIÓN DE CONTRATO, AGRAVIO O DE OTRO MODO, QUE SURJA DE, FUERA DE O EN 
CONEXIÓN CON EL SOFTWARE O EL USO U OTRO TIPO DE ACCIONES EN EL SOFTWARE.
"""

import os
import subprocess
import sys

# Verificar e instalar paquetes
def check_and_install_packages():
    required_packages = ["python-docx", "pillow", "google-cloud-vision", "openpyxl"]
    for package in required_packages:
        try:
            __import__(package if package != "python-docx" else "docx")
        except ImportError:
            print(f"Instalando {package}...")
            subprocess.check_call([sys.executable, "-m", "pip", "install", package])

check_and_install_packages()

import re
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import Workbook
from PIL import Image
from google.cloud import vision


def clean_filename(filename):
    """Limpia el nombre del archivo para mostrar solo el nombre sin extensión y con la primera letra en mayúscula."""
    name = os.path.splitext(filename)[0]  # Remover la extensión
    name = re.sub(r'[_]', ' ', name)      # Reemplazar guiones bajos con espacios
    return name.capitalize()  

def add_header(doc, location, name, dni, ref_expediente):
    """Agrega el título y encabezado de documento con la información dada."""
    title = f"LISTADO DE ELECTRODOMÉSTICOS, MUEBLES Y ENSERES DAÑADOS 29/10/2024\nVIVIENDA DE {location}\n"
    sub_title = f"({name}\nD.N.I. {dni})\n\nS/Ref. Expediente {ref_expediente}"

    # Crear el título principal
    title_para = doc.add_paragraph(title)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title_para.runs:
        run.font.size = Pt(16)           # Establecer tamaño de fuente a 16 puntos
        run.font.bold = True             # Asegurar que esté en negrita
        run.font.color.rgb = RGBColor(0, 0, 0)  # Asegurar color negro

    # Crear el subtítulo
    sub_title_para = doc.add_paragraph(sub_title)
    sub_title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in sub_title_para.runs:
        run.font.size = Pt(12)           # Tamaño de fuente menor para el subtítulo
        run.font.bold = False            # Subtítulo sin negrita
        run.font.color.rgb = RGBColor(0, 0, 0)  # Asegurar color negroR

def add_page_numbers(doc):
    """Añade números de página al documento."""
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Campo de número de página
    page_num_field = footer.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    page_num_field._r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    page_num_field._r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    page_num_field._r.append(fldChar)

def add_table_of_contents(doc):
    """Inserta un marcador de posición para una tabla de contenidos (ToC) con hipervínculos."""
    # Crear un párrafo para la ToC
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    
    # Crear el campo de la ToC
    fldChar1 = OxmlElement('w:fldChar')  # Crear elemento de inicio de campo
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')  # Instrucciones para el campo
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # Nivel de encabezado y opciones de ToC (hipervínculos)
    
    fldChar2 = OxmlElement('w:fldChar')  # Crear elemento de separación de campo
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')  # Crear elemento de fin de campo
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    # Añadir los elementos al run del párrafo
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centramos la ToC


def add_images_to_doc(doc, folder_path, toc_data, base_depth, page_counter, page_width_inches=8.27, margin_inches=1.0, scale_factor=0.5):
    """Agrega las imágenes de una carpeta al documento y genera datos para la tabla de contenidos."""
    max_image_width = (page_width_inches - 2 * margin_inches) * scale_factor * 96

    for filename in os.listdir(folder_path):
        if filename.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.gif', '.jfif')):
            clean_name = clean_filename(filename)
            toc_data.append((clean_name, len(doc.paragraphs) + 1, "image"))  # Sección, página, precio placeholder

            # Añadir el nombre de la imagen como un subtítulo (Heading, ajustado al nivel de la carpeta)
            para = doc.add_paragraph(clean_name, style=f'Heading {base_depth + 2}')  # Ajuste dinámico del nivel
            for run in para.runs:
                run.font.bold = False
                run.font.italic = False
                run.font.color.rgb = RGBColor(0, 0, 0)  # Establecer color negro
            
            image = Image.open(os.path.join(folder_path, filename))

             # Convertir a modo RGB si tiene un canal alfa
            if image.mode in ('RGBA', 'P'):
                image = image.convert('RGB')

            width, height = image.size
            aspect_ratio = height / width

            # Ajustar el tamaño manteniendo la proporción
            new_width = int(min(max_image_width, width))  # Asegura que no agrandamos la imagen original
            new_height = int(new_width * aspect_ratio)
            image = image.resize((new_width, new_height), Image.LANCZOS)

            temp_image_path = "temp_resized_image.jpg"
            image.save(temp_image_path, format = "JPEG", quality=95)

           
            
            # Insertar la imagen en el documento
            doc.add_picture(temp_image_path, width=Inches(new_width / 96))
            os.remove(temp_image_path)
            
            # Incrementar el contador de páginas después de cada imagen y añadir salto de página
            page_counter += 1
            doc.add_paragraph().add_run().add_break()

    return page_counter
    
def label_image(image_path):
    """Usa Google Vision para obtener etiquetas de una imagen."""
    client = vision.ImageAnnotatorClient()
    with io.open(image_path, 'rb') as image_file:
        content = image_file.read()

    image = vision.Image(content=content)
    response = client.label_detection(image=image)
    labels = response.label_annotations

    if labels:
        return labels[0].description
    return None

def rename_images_in_folder(folder_path):
    """Renombra las imágenes en la carpeta usando etiquetas de Google Vision."""
    for root, dirs, files in os.walk(folder_path):
        for filename in files:
            if filename.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.gif', '.jfif')):
                image_path = os.path.join(root, filename)
                label = label_image(image_path)
                
                if label:
                    new_filename = f"{label}_{filename}"
                    new_image_path = os.path.join(root, new_filename)
                    os.rename(image_path, new_image_path)
                    print(f"{filename} renombrado a {new_filename}")

def crear_inventario(root_folder, output_doc, location, name, dni, ref_expediente, scale_factor = 0.5,
                     rename_images = False, create_doc=False):
    """Crea un documento de inventario con las mejoras solicitadas, según las opciones especificadas."""
    if rename_images:
        rename_images_in_folder(root_folder)

    if create_doc:
        doc = Document()
        add_header(doc, location, name, dni, ref_expediente)
        add_page_numbers(doc)
        
        # Insertar la tabla de contenidos al inicio
        add_table_of_contents(doc)

        # Inicializar datos para la tabla de contenidos
        toc_data = []

        page_counter = 1  # Comenzamos en la primera página después del encabezado
        # Obtener el nivel base de la carpeta raíz para ajustar los encabezados
        base_depth = root_folder.rstrip(os.path.sep).count(os.path.sep)

        # Agregar tabla de contenidos vacía (se llenará después)
        toc_start_paragraph = doc.add_paragraph("Tabla de Contenidos")
        toc_start_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        toc_start_index = len(doc.paragraphs)

        # Agregar imágenes y nombres de carpetas al documento
        for root, dirs, files in os.walk(root_folder):
            if files:
                folder_name = os.path.basename(root)
                # Calcular el nivel de encabezado basado en la profundidad de la carpeta
                current_depth = root.count(os.path.sep) - base_depth + 1
                folder_para = doc.add_paragraph(folder_name, style=f'Heading {current_depth}')
                folder_para.runs[0].font.bold = True
                toc_data.append((folder_name, page_counter, "section"))
                page_counter += 1
                add_images_to_doc(doc, root, toc_data, current_depth, page_counter, scale_factor = scale_factor)

        # Agregar tabla de contenidos después del encabezado
        #add_table_of_contents(doc, toc_data)

        # Agregar secciones finales
        doc.add_paragraph("Pintura", style='Heading 1').runs[0].font.bold = True
        toc_data.append(("Pintura", page_counter, "section"))
        doc.add_paragraph("Daños estructura", style='Heading 1').runs[0].font.bold = True
        toc_data.append(("Daños estructura", page_counter, "section"))
        
        doc.save(f"../documentos/{output_doc}.docx")
        print(f"Documento de inventario guardado en: pynventario/documentos/{output_doc}.docx")

        # Generar archivo Excel de la tabla de contenidos
        generate_excel_toc(toc_data, "inventario_toc")

def generate_excel_toc(toc_data, output_xlsx):
    """Genera un archivo Excel con la tabla de contenidos y columnas adicionales."""
    wb = Workbook()
    ws = wb.active
    ws.append(["Objeto", "Página listado", "Precio"])

    for entry in toc_data:
        section, page, entry_type = entry
        ws.append([section, None, None])
        # Poner en negrita solo las secciones principales
        if entry_type == "section":
            ws[f"A{ws.max_row}"].font = ws[f"A{ws.max_row}"].font.copy(bold=True)

    wb.save(f"../documentos/{output_xlsx}.xlsx")
    print(f"Tabla de contenidos guardada en: pynventario/documentos/{output_xlsx}.xlsx")

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Generador de inventario de imágenes y tabla de contenidos.")
    parser.add_argument("root_folder", type=str, help="Ruta de la carpeta raíz que contiene las imágenes.")
    parser.add_argument("output_doc", type=str, help="Nombre del archivo de inventario de salida.")
    parser.add_argument("location", type=str, help="Nombre de la población.")
    parser.add_argument("name", type=str, help="Nombre de la persona.")
    parser.add_argument("dni", type=str, help="DNI de la persona.")
    parser.add_argument("ref_expediente", type=str, help="Referencia del expediente.")
    parser.add_argument("scale_factor", type=float, help="Fracción del ancho de página para la foto (p.ej: 0.5 es la mitad del ancho).")
    parser.add_argument("--rename_images", action="store_true", help="Renombra las imágenes usando Google Vision.")
    parser.add_argument("--create_doc", action="store_true", help="Genera el documento de inventario.")

    args = parser.parse_args()

    crear_inventario(
        root_folder=args.root_folder,
        output_doc=args.output_doc,
        location=args.location,
        name=args.name,
        dni=args.dni,
        ref_expediente=args.ref_expediente,
        rename_images=args.rename_images,
        create_doc=args.create_doc,
        scale_factor = scale_factor
    )
